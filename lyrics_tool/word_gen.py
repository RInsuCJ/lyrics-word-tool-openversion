"""
word_gen.py — Generate a printable Word (.docx) document from song lyrics.

Auto-layout algorithm:
  1. Count total lines in the lyrics.
  2. Scan font sizes from 16 pt down to 8 pt (0.5 pt steps).
     Return the LARGEST size whose estimated content height fits on one A4 page.
  3. If no single-column size fits, switch to two-column layout
     (title spans full width; lyrics split into two columns).

Spacing model (used in both estimation and generation):
  - Title paragraph : space_before=4pt, EXACTLY title_pt*1.25, space_after=6pt
  - Lyric paragraph : space_before=0,   EXACTLY body_pt*1.30,  space_after=1pt
  - No blank line between title and lyrics (direct connection).
"""

import math
import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Page geometry ─────────────────────────────────────────────────────────────
PAGE_W_CM     = 21.0
PAGE_H_CM     = 29.7
MARGIN_TOP    = 2.54
MARGIN_BOTTOM = 2.54
MARGIN_LEFT   = 3.18
MARGIN_RIGHT  = 3.18

CM_TO_PT = 72 / 2.54          # 1 cm ≈ 28.346 pt

USABLE_H_PT = (PAGE_H_CM - MARGIN_TOP - MARGIN_BOTTOM) * CM_TO_PT   # ≈ 698 pt
USABLE_W_PT = (PAGE_W_CM - MARGIN_LEFT - MARGIN_RIGHT) * CM_TO_PT   # ≈ 415 pt

# ── Typography / spacing constants ────────────────────────────────────────────
DEFAULT_FONT     = "Arial Unicode MS"
DEFAULT_TITLE_PT = 22.0

BODY_MIN  = 8.0
BODY_MAX  = 16.0
BODY_STEP = 0.5

# Title spacing (applied to title paragraph in generate_word)
TITLE_SPACE_BEFORE_PT = 4.0
TITLE_LINE_FACTOR     = 1.25   # EXACTLY title_pt * 1.25
TITLE_SPACE_AFTER_PT  = 6.0

# Body/lyric line spacing (applied to every lyric paragraph)
# EXACTLY spacing: actual Word height = font_pt * BODY_LINE_FACTOR
# Measured on Arial Unicode MS (CJK): actual ≈ 1.32x font size.
# Using 1.35 gives a small safety buffer above the measured value.
BODY_LINE_FACTOR     = 1.35
BODY_SPACE_AFTER_PT  = 1.0

# Safety margin: only fill up to 96 % of the usable height.
SAFETY = 0.96

# Two-column gap
TWO_COL_GAP_CM = 0.8

# Extra overhead added to two-column height estimates.
TWO_COL_TABLE_OVERHEAD_PT = 25.0

# ── Pinyin / romanisation spacing ─────────────────────────────────────────────
# When with_pinyin=True each lyric line is preceded by a romanisation line.
# Pinyin uses a smaller font (PINYIN_FONT_RATIO × body_pt) with its own
# EXACTLY line spacing.  No space_after — it sits directly above the lyric.
PINYIN_FONT_RATIO  = 0.70   # pinyin_pt = body_pt * 0.70
PINYIN_LINE_FACTOR = 1.20   # EXACTLY height for pinyin line
PINYIN_COLOR_HEX   = "808080"  # grey colour for pinyin text in Word
BLANK_LINE_PT      = 12.0      # fixed height for blank lines between verses



# ── Height estimators ─────────────────────────────────────────────────────────

def _title_h(title_pt: float, n_lines: int = 1) -> float:
    """Height of the title block in points.  n_lines = 1 or 2."""
    return (TITLE_SPACE_BEFORE_PT
            + title_pt * TITLE_LINE_FACTOR * n_lines
            + TITLE_SPACE_AFTER_PT)


def _est_text_width(text: str, font_pt: float) -> float:
    """
    Rough width estimate in points for Arial Unicode MS.
    CJK characters are 1 em wide; Latin/ASCII ≈ 0.55 em.
    """
    w = 0.0
    for ch in text:
        cp = ord(ch)
        if (0x4E00 <= cp <= 0x9FFF or   # CJK unified
                0x3400 <= cp <= 0x4DBF or   # CJK extension A
                0xF900 <= cp <= 0xFAFF or   # CJK compatibility
                0x3000 <= cp <= 0x303F or   # CJK symbols/punctuation
                0xFF00 <= cp <= 0xFFEF):    # fullwidth Latin
            w += font_pt
        else:
            w += font_pt * 0.55             # Latin, digits, spaces
    return w


def _resolve_title(title: str, artist: str,
                   title_pt: float) -> tuple[list[str], float]:
    """
    Decide how to render the heading.

    Returns (lines, effective_pt) where:
      lines = ['《title》artist']           →  one centred paragraph
      lines = ['《title》', 'artist']       →  two centred paragraphs

    Priority:
      1. «title» artist on ONE line  at current font size
      2. «title» on line 1 / artist on line 2  at current font size
      3. Repeat 1→2 reducing font by 0.5 pt until ≥ 14 pt
      4. Absolute fallback: two lines at 14 pt
    """
    one   = f"\u300a{title}\u300b{artist}"  # 《title》artist
    line1 = f"\u300a{title}\u300b"           # 《title》
    line2 = artist

    for pt in (title_pt - step * 0.5 for step in range(9)):  # try down to -4 pt
        pt = max(14.0, round(pt, 1))
        if _est_text_width(one, pt) <= USABLE_W_PT:
            return [one], pt
        if (_est_text_width(line1, pt) <= USABLE_W_PT and
                _est_text_width(line2, pt) <= USABLE_W_PT):
            return [line1, line2], pt
        if pt <= 14.0:
            break
    return [line1, line2], 14.0


def _body_line_h(body_pt: float) -> float:
    """Height of one body/lyric line in points."""
    return body_pt * BODY_LINE_FACTOR + BODY_SPACE_AFTER_PT


def _pinyin_line_h(body_pt: float) -> float:
    """Height of one pinyin annotation line (directly above a lyric line)."""
    return body_pt * PINYIN_FONT_RATIO * PINYIN_LINE_FACTOR  # no space_after


def _unit_h(body_pt: float, with_pinyin: bool = False) -> float:
    """Height of one lyric 'unit' (pinyin row if enabled + lyric row)."""
    h = _body_line_h(body_pt)
    if with_pinyin:
        h += _pinyin_line_h(body_pt)
    return h


def _estimate_single(n_lines: int, body_pt: float, title_pt: float,
                     with_pinyin: bool = False,
                     n_blank: int = 0) -> float:
    """Estimate total content height for single-column layout.
    
    n_blank: number of blank lines in the lyrics (verse separators).
    When with_pinyin, blank lines are only BLANK_LINE_PT tall instead of full unit_h.
    """
    n_content = n_lines - n_blank
    h = _title_h(title_pt) + _unit_h(body_pt, with_pinyin) * n_content
    if with_pinyin:
        h += BLANK_LINE_PT * n_blank
    else:
        h += _body_line_h(body_pt) * n_blank
    return h


def _estimate_two_col(n_lines: int, body_pt: float, title_pt: float,
                      with_pinyin: bool = False,
                      n_blank: int = 0) -> float:
    """
    Estimate content height for two-column layout (includes table overhead).
    Title spans full width; lyrics split into two columns inside a table.
    """
    col_lines = math.ceil(n_lines / 2)
    col_blank = math.ceil(n_blank / 2)  # blanks split roughly evenly
    col_content = col_lines - col_blank
    h = _title_h(title_pt) + _unit_h(body_pt, with_pinyin) * col_content
    if with_pinyin:
        h += BLANK_LINE_PT * col_blank
    else:
        h += _body_line_h(body_pt) * col_blank
    return h + TWO_COL_TABLE_OVERHEAD_PT


def find_optimal_layout(lyrics: str,
                        title_pt: float = DEFAULT_TITLE_PT,
                        max_body_pt: float = BODY_MAX,
                        with_pinyin: bool = False,
                        ) -> tuple[float, bool]:
    """
    Return (body_size_pt, use_two_columns).

    Layout priority (largest font wins):
      1. Single column, fits in 1 page
      2. Two columns,   fits in 1 page
      3. Two columns,   fits in 2 pages
      4. Absolute fallback: BODY_MIN, two columns
    """
    n_lines   = len(lyrics.splitlines())
    n_blank   = sum(1 for l in lyrics.splitlines() if not l.strip())
    cap       = min(max_body_pt, BODY_MAX)
    steps     = round((cap - BODY_MIN) / BODY_STEP) + 1
    target_1p = USABLE_H_PT * SAFETY
    target_2p = 2 * USABLE_H_PT * SAFETY

    # Phase 1: single column, 1 page (always tried regardless of pinyin)
    for i in range(steps):
        size = round(cap - i * BODY_STEP, 1)
        if _estimate_single(n_lines, size, title_pt, with_pinyin, n_blank) <= target_1p:
            return (size, False)

    if with_pinyin:
        # Pinyin lines are Latin text — they overflow in narrow two-column cells.
        # When pinyin is on, stay single-column and allow up to 2 pages instead.
        for i in range(steps):
            size = round(cap - i * BODY_STEP, 1)
            if _estimate_single(n_lines, size, title_pt, with_pinyin, n_blank) <= target_2p:
                return (size, False)
        return (BODY_MIN, False)   # fallback: minimum font, single column
    else:
        # No pinyin — original two-column logic unchanged
        for i in range(steps):
            size = round(cap - i * BODY_STEP, 1)
            if _estimate_two_col(n_lines, size, title_pt, with_pinyin, n_blank) <= target_1p:
                return (size, True)

        for i in range(steps):
            size = round(cap - i * BODY_STEP, 1)
            if _estimate_two_col(n_lines, size, title_pt, with_pinyin, n_blank) <= target_2p:
                return (size, True)

        return (BODY_MIN, True)   # absolute fallback


# ── Word XML helpers ──────────────────────────────────────────────────────────

def _set_run_font(run, font_name: str, size_pt: float | None = None,
                  bold: bool = False):
    """Apply font to a run, covering Latin, CJK and complex-script slots."""
    run.font.name = font_name
    rPr    = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"),       font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold:
        run.font.bold = True


def _set_exact_spacing(para,
                        line_pt: float,
                        space_before_pt: float = 0.0,
                        space_after_pt:  float = BODY_SPACE_AFTER_PT):
    """
    Apply EXACTLY line spacing so that each paragraph's height is predictable.
    This removes any font-metric variation (important for CJK fonts).
    """
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing      = Pt(line_pt)
    fmt.space_before      = Pt(space_before_pt)
    fmt.space_after       = Pt(space_after_pt)


def _add_lyric_paragraph(doc: Document, text: str,
                          font_name: str, body_pt: float) -> "Paragraph":
    """Add one lyric line with compact, predictable spacing."""
    para = doc.add_paragraph(style="Normal")
    if text:
        run = para.add_run(text)
        _set_run_font(run, font_name, body_pt)
    _set_exact_spacing(
        para,
        line_pt        = body_pt * BODY_LINE_FACTOR,
        space_before_pt = 0.0,
        space_after_pt  = BODY_SPACE_AFTER_PT,
    )
    return para


def _cm_to_twips(cm: float) -> int:
    """Convert centimetres to twips (1 pt = 20 twips, 1 cm = 28.346 pt)."""
    return round(cm * 1440 / 2.54)


# ── Two-column table helpers ─────────────────────────────────────────────────

TWO_COL_GAP_CM = 0.8   # visual gap between columns (via cell right-padding)


def _clear_all_borders(table):
    """Remove all visible borders from a table and its cells."""
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Table-level borders
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # Cell-level borders
    for cell in table.rows[0].cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   "none")
            el.set(qn("w:sz"),    "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
            tcBorders.append(el)
        tcPr.append(tcBorders)



def _set_cell_width(cell, width_twips: int):
    """Set a table cell to an exact width in twips."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_cell_margins(cell,
                      top: float = 0, bottom: float = 0,
                      left: float = 0, right: float = 0):
    """Set cell padding (in pt)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(int(val * 20)))   # pt → twips (1 pt = 20 twips)
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _fill_cell(cell, lines: list[str], font_name: str, body_pt: float,
               annotated_col: list[tuple[str, str]] | None = None,
               pinyin_swapped: bool = False):
    """Fill a table cell with lyric lines (optionally with pinyin rows above)."""
    source = annotated_col if annotated_col is not None else [(None, l) for l in lines]
    first_para = True
    for rom, line in source:
        # Pinyin row (if enabled)
        if rom is not None:   # rom is None only when annotated_col not provided
            if not line.strip():
                # Blank line = verse separator — single 12pt line, no pinyin
                if first_para:
                    para = cell.paragraphs[0]
                    first_para = False
                else:
                    para = cell.add_paragraph()
                _set_exact_spacing(para,
                                   line_pt=BLANK_LINE_PT,
                                   space_before_pt=0, space_after_pt=0)
                continue

            if pinyin_swapped:
                # Swapped: lyrics (small grey) on top, pinyin (big black) below
                # Top row: lyrics as small grey annotation
                if first_para:
                    sub_para = cell.paragraphs[0]
                    first_para = False
                else:
                    sub_para = cell.add_paragraph()
                pinyin_pt = body_pt * PINYIN_FONT_RATIO
                sub_run = sub_para.add_run(line if line else ' ')
                _set_run_font(sub_run, font_name, pinyin_pt, bold=False)
                rPr = sub_run._r.get_or_add_rPr()
                color_el = OxmlElement('w:color')
                color_el.set(qn('w:val'), PINYIN_COLOR_HEX)
                rPr.append(color_el)
                _set_exact_spacing(sub_para,
                                   line_pt         = pinyin_pt * PINYIN_LINE_FACTOR,
                                   space_before_pt = 0,
                                   space_after_pt  = 0)
                # Bottom row: pinyin as main text
                main_para = cell.add_paragraph()
                if rom:
                    run = main_para.add_run(rom)
                    _set_run_font(run, font_name, body_pt)
                _set_exact_spacing(
                    main_para,
                    line_pt         = body_pt * BODY_LINE_FACTOR,
                    space_before_pt = 0.0,
                    space_after_pt  = BODY_SPACE_AFTER_PT,
                )
            else:
                # Normal: pinyin (small grey) on top, lyrics (big black) below
                if first_para:
                    pin_para = cell.paragraphs[0]
                    first_para = False
                else:
                    pin_para = cell.add_paragraph()
                pin_run = pin_para.add_run(rom if rom else ' ')
                pinyin_pt = body_pt * PINYIN_FONT_RATIO
                _set_run_font(pin_run, font_name, pinyin_pt, bold=False)
                rPr = pin_run._r.get_or_add_rPr()
                color_el = OxmlElement('w:color')
                color_el.set(qn('w:val'), PINYIN_COLOR_HEX)
                rPr.append(color_el)
                _set_exact_spacing(pin_para,
                                   line_pt         = pinyin_pt * PINYIN_LINE_FACTOR,
                                   space_before_pt = 0,
                                   space_after_pt  = 0)
                # Lyric row always gets its own new paragraph
                para = cell.add_paragraph()
                if line:
                    run = para.add_run(line)
                    _set_run_font(run, font_name, body_pt)
                _set_exact_spacing(
                    para,
                    line_pt         = body_pt * BODY_LINE_FACTOR,
                    space_before_pt = 0.0,
                    space_after_pt  = BODY_SPACE_AFTER_PT,
                )
        else:
            # No pinyin — original behaviour
            if first_para:
                para = cell.paragraphs[0]
                first_para = False
            else:
                para = cell.add_paragraph()

            if line:
                run = para.add_run(line)
                _set_run_font(run, font_name, body_pt)
            _set_exact_spacing(
                para,
                line_pt         = body_pt * BODY_LINE_FACTOR,
                space_before_pt = 0.0,
                space_after_pt  = BODY_SPACE_AFTER_PT,
            )


def _make_two_col_table(doc: Document, lyrics_lines: list[str],
                         font_name: str, body_pt: float,
                         annotated: list[tuple[str, str]] | None = None,
                         pinyin_swapped: bool = False):
    """
    Create a borderless 2-column table and fill it with lyrics.
    Left column gets the first half of lines; right column gets the rest.
    If annotated is provided, each cell also renders pinyin rows.
    """
    n    = len(lyrics_lines)
    half = math.ceil(n / 2)
    col1 = lyrics_lines[:half]
    col2 = lyrics_lines[half:]
    ann1 = annotated[:half]  if annotated else None
    ann2 = annotated[half:]  if annotated else None

    table = doc.add_table(rows=1, cols=2)
    _clear_all_borders(table)

    # Total usable width in twips
    usable_twips = _cm_to_twips(PAGE_W_CM - MARGIN_LEFT - MARGIN_RIGHT)
    gap_twips    = _cm_to_twips(TWO_COL_GAP_CM)
    col_w        = (usable_twips - gap_twips) // 2

    left_cell, right_cell = table.rows[0].cells

    _set_cell_width(left_cell,  col_w)
    _set_cell_width(right_cell, col_w)

    # Right padding on left cell creates the visual gap
    gap_pt = TWO_COL_GAP_CM * CM_TO_PT
    _set_cell_margins(left_cell,  top=0, bottom=0, left=0, right=gap_pt / 2)
    _set_cell_margins(right_cell, top=0, bottom=0, left=gap_pt / 2, right=0)

    _fill_cell(left_cell,  col1, font_name, body_pt, annotated_col=ann1, pinyin_swapped=pinyin_swapped)
    _fill_cell(right_cell, col2, font_name, body_pt, annotated_col=ann2, pinyin_swapped=pinyin_swapped)

    return table


def _reset_doc_defaults(doc: Document):
    """
    Zero out the document's default paragraph spacing (docDefaults/pPrDefault).

    python-docx's built-in template has:
        w:spacing w:after="200"  (10 pt space-after)
        w:spacing w:line="276"   (1.15x auto line spacing)

    These defaults bleed into table-cell paragraphs even when EXACTLY spacing
    is applied as a direct override, causing Word to add extra height that our
    estimator cannot predict.  Zeroing them here ensures EXACTLY truly means
    the exact value we set, throughout the document.
    """
    styles_part = doc.part.styles
    styles_el   = styles_part._element           # <w:styles>

    docDef = styles_el.find(qn("w:docDefaults"))
    if docDef is None:
        return

    pPrDef = docDef.find(qn("w:pPrDefault"))
    if pPrDef is None:
        return

    pPr = pPrDef.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        pPrDef.append(pPr)

    # Remove any existing spacing element and replace with zeroed values
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)

    sp = OxmlElement("w:spacing")
    sp.set(qn("w:after"),    "0")    # no space-after
    sp.set(qn("w:line"),     "240")  # single = 240 twips (1.0x), not 1.15x
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


# ── Main public function ───────────────────────────────────────────────────────

def _add_pinyin_paragraph(container, rom: str, font_name: str,
                           body_pt: float) -> None:
    """Add a small grey romanisation paragraph directly above a lyric line."""
    pinyin_pt = body_pt * PINYIN_FONT_RATIO
    para = container.add_paragraph()
    para.style = container.styles["Normal"] if hasattr(container, 'styles') \
                 else container._body._parent.styles["Normal"]
    run = para.add_run(rom if rom else ' ')
    _set_run_font(run, font_name, pinyin_pt, bold=False)
    # Grey colour
    rPr = run._r.get_or_add_rPr()
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), PINYIN_COLOR_HEX)
    rPr.append(color_el)
    _set_exact_spacing(
        para,
        line_pt         = pinyin_pt * PINYIN_LINE_FACTOR,
        space_before_pt = 0,
        space_after_pt  = 0,   # no gap — lyric sits immediately below
    )


def generate_word(
    output_path: str,
    title: str,
    artist: str,
    lyrics: str,
    font_name: str = DEFAULT_FONT,
    title_size: float = DEFAULT_TITLE_PT,
    add_page_break_at_end: bool = False,
    force_body_pt: float | None = None,
    force_two_cols: bool | None = None,
    with_pinyin: bool = False,
    pinyin_lang: str = 'cantonese',
    pinyin_swapped: bool = False,
) -> tuple[str, float, bool]:
    """
    Generate a Word document with automatic layout.
    Returns (output_path, body_size_used, used_two_cols).

    If force_body_pt / force_two_cols are given, they override the auto-layout
    calculation (used when the user manually adjusts sizes in the preview dialog).
    """
    # ── 1. Determine layout ──────────────────────────────────────────────────
    if force_body_pt is not None and force_two_cols is not None:
        body_pt   = force_body_pt
        two_cols  = force_two_cols
    else:
        body_pt, two_cols = find_optimal_layout(
            lyrics, title_pt=title_size, with_pinyin=with_pinyin)
        if force_body_pt  is not None: body_pt  = force_body_pt
        if force_two_cols is not None: two_cols = force_two_cols

    # ── 1b. Pre-compute romanisation if needed ───────────────────────────────
    annotated: list[tuple[str, str]] | None = None
    if with_pinyin:
        from lyrics_tool.pinyin import annotate_lyrics
        annotated = annotate_lyrics(lyrics, lang=pinyin_lang)

    # ── 2. Create document ───────────────────────────────────────────────────
    doc     = Document()

    # Zero out docDefaults paragraph spacing.
    # The template has w:after="200" (10pt) and w:line="276" (1.15x auto).
    # These defaults bleed into table-cell paragraphs even when EXACTLY
    # spacing is applied as a direct override, causing Word to render
    # cell content taller than estimated.
    _reset_doc_defaults(doc)

    section = doc.sections[0]
    section.page_width    = Cm(PAGE_W_CM)
    section.page_height   = Cm(PAGE_H_CM)
    section.top_margin    = Cm(MARGIN_TOP)
    section.bottom_margin = Cm(MARGIN_BOTTOM)
    section.left_margin   = Cm(MARGIN_LEFT)
    section.right_margin  = Cm(MARGIN_RIGHT)

    # ── 3. Title ─────────────────────────────────────────────────────────────
    # Resolve 1-line vs 2-line layout (with possible font reduction).
    title_lines, eff_title_pt = _resolve_title(title, artist, title_size)
    n_title_lines = len(title_lines)

    for idx, line_text in enumerate(title_lines):
        is_first = idx == 0
        is_last  = idx == n_title_lines - 1

        para = doc.add_paragraph()
        para.style     = doc.styles["Normal"]
        para.alignment = 1  # CENTER

        run = para.add_run(line_text)
        _set_run_font(run, font_name, eff_title_pt, bold=True)
        _set_exact_spacing(
            para,
            line_pt         = eff_title_pt * TITLE_LINE_FACTOR,
            space_before_pt = TITLE_SPACE_BEFORE_PT if is_first else 0,
            space_after_pt  = TITLE_SPACE_AFTER_PT  if is_last  else 0,
        )

    # ── 4. Lyrics ─────────────────────────────────────────────────────────────
    lines = lyrics.splitlines()

    if two_cols:
        # For two-column layout pass the annotated pairs so the table builder
        # can interleave pinyin rows inside each cell.
        _make_two_col_table(doc, lines, font_name, body_pt,
                            annotated=annotated,
                            pinyin_swapped=pinyin_swapped)
    else:
        if annotated:
            for rom, line in annotated:
                if not line.strip():
                    # Blank line = verse separator — single 12pt line, no pinyin
                    para = doc.add_paragraph(style="Normal")
                    _set_exact_spacing(para,
                                       line_pt=BLANK_LINE_PT,
                                       space_before_pt=0, space_after_pt=0)
                elif pinyin_swapped:
                    # Swapped: lyrics (small grey) on top, pinyin (big black) below
                    _add_pinyin_paragraph(doc, line, font_name, body_pt)
                    _add_lyric_paragraph(doc, rom, font_name, body_pt)
                else:
                    _add_pinyin_paragraph(doc, rom, font_name, body_pt)
                    _add_lyric_paragraph(doc, line, font_name, body_pt)
        else:
            for line in lines:
                _add_lyric_paragraph(doc, line, font_name, body_pt)

    # ── 5. Optional page break ────────────────────────────────────────────────
    if add_page_break_at_end:
        p   = doc.add_paragraph()
        run = p.add_run()
        br  = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    # ── 6. Save ───────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path, body_pt, two_cols
